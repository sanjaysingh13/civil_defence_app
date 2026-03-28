"""
Management command: seed_vehicles

Usage:
    uv run python manage.py seed_vehicles
    uv run python manage.py seed_vehicles --docx /path/to/CDRV Details.docx
    uv run python manage.py seed_vehicles --dry-run

What it does:
    1. Reads "CDRV Details.docx" — a Word table with columns:
           Sl. No | Allotted District | Vehicle No | Vehicle type
    2. For every row (skipping the header):
         - Maps the district name to a DB Unit using UNIT_NAME_MAP.
         - Units that don't yet exist are created automatically.
         - Creates / updates a Vehicle row keyed on registration_no.
    3. Vehicle type is mapped:
         "Big CDRV"  → VehicleType.BIG_CDRV
         "Mini CDRV" → VehicleType.MINI_CDRV
         anything else → VehicleType.OTHER  (with a warning)
    4. Default status = AVAILABLE for all imported vehicles.

Known data issues (handled automatically):
    • Kolkata: 8 vehicles in the docx — skipped with a warning (same as the
      equipment import; the 3 Kolkata DB units cannot be auto-split).
    • Duplicate registration numbers in the source file:
        - WB-37C-6770 appears for both Jalpaiguri (R9) and Alipurduar (R52).
        - WB-41G-2767 appears for both Darjeeling (R4) and Kalimpong (R50).
      The FIRST occurrence is kept; any later duplicate is skipped with a warning.

Idempotency:
    Vehicles are upserted on registration_no (update_or_create), so the
    command is safe to re-run — it updates existing rows rather than creating
    duplicates.  The Kolkata skip logic and duplicate guard remain active on
    every run.

Unit name mapping:
    The docx uses mixed-case district names.  UNIT_NAME_MAP translates each
    to the canonical name stored in the DB (preserving known typos such as
    "KALIMPOLNG").  New units not yet in the DB are created on first run.
"""

from pathlib import Path

from django.core.management.base import BaseCommand

from civil_defence_app.fleet.models import Vehicle, VehicleStatus, VehicleType
from civil_defence_app.personnel.models import Unit


# ─────────────────────────────────────────────────────────────────────────────
# DEFAULT FILE PATH
#
# This file lives at:
#   civil_defence_app/civil_defence_app/fleet/management/commands/seed_vehicles.py
# parents[5] therefore resolves to the project root:  civil_defence/
# ─────────────────────────────────────────────────────────────────────────────

DEFAULT_DOCX_PATH = Path(__file__).resolve().parents[5] / "CDRV Details.docx"


# ─────────────────────────────────────────────────────────────────────────────
# UNIT NAME MAP
#
# Key   = district name as it appears in the "Allotted District" column of the
#         Word table (strip-normalised to single spaces before lookup).
# Value = canonical unit name stored in the DB.
#         None → skip this row entirely (used for Kolkata and future unknowns).
#
# Rules:
#   • "Cooch Bihar" in the docx is "Cooch Behar" / "COOCHBEHAR" in the DB.
#   • "Bardhaman" districts are stored under the older "Burdwan" spelling.
#   • "KALIMPOLNG" — DB has this typo; we preserve it rather than fix it so
#     this command doesn't break if someone else also queries by that name.
#   • Kolkata → None: 8 vehicles listed under a single row, but 3 separate DB
#     units (CSA, NSA, SSA) — cannot be split automatically.
# ─────────────────────────────────────────────────────────────────────────────

UNIT_NAME_MAP: dict[str, str | None] = {
    "Alipurduar":       "ALIPURDUAR",
    "Bankura":          "BANKURA",
    "Birbhum":          "BIRBHUM",
    "Cooch Bihar":      "COOCHBEHAR",       # docx spelling differs from DB
    "Dakshin Dinajpur": "DAKSHIN DINAJPUR",
    "Darjeeling":       "DARJEELING",
    "Hooghly":          "HOOGHLY",
    "Howrah":           "HOWRAH",
    "Jalpaiguri":       "JALPAIGURI",
    "Jhargram":         "JHARGRAM",
    "Kalimpong":        "KALIMPOLNG",        # DB has this typo — preserve it
    "Malda":            "MALDA",
    "Murshidabad":      "MURSHIDABAD",
    "Nadia":            "NADIA",
    "North 24 Pgs":     "NORTH24 PGS",
    "Paschim Bardhaman":"PASCHIM BURDWAN",   # Bardhaman = Burdwan (older spelling)
    "Paschim Medinipur":"PASCHIM MEDINIPUR",
    "Purba Bardhaman":  "PURBA BURDWAN",
    "Purba Medinipur":  "PURBA MEDINIPUR",   # created by seed_equipment if not present
    "Purulia":          "PURULIA",           # created by seed_equipment if not present
    "South 24 Pgs":     "SOUTH24 PGS",
    "Uttar Dinajpur":   "UTTAR DINAJ PUR",
    # Kolkata → None: a single docx row covers 3 DB units (CSA/NSA/SSA).
    # Cannot be split automatically; skipped with a warning.
    "Kolkata":          None,
}


# ─────────────────────────────────────────────────────────────────────────────
# VEHICLE TYPE MAP
#
# Translates the "Vehicle type" column in the docx to a VehicleType choice.
# ─────────────────────────────────────────────────────────────────────────────

VEHICLE_TYPE_MAP: dict[str, str] = {
    "Big CDRV":  VehicleType.BIG_CDRV,
    "Mini CDRV": VehicleType.MINI_CDRV,
}


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def normalise(text: str) -> str:
    """
    Strip leading/trailing whitespace and collapse internal runs of whitespace
    (spaces, non-breaking spaces, tabs) down to a single space.

    Word documents sometimes have non-breaking spaces (\\xa0) or extra spaces
    in table cells; this helper makes sure lookups in UNIT_NAME_MAP are robust
    to those variations.
    """
    # ' '.join(text.split()) handles all whitespace variants including \\xa0
    return " ".join(text.split())


def parse_vehicle_table(docx_path: Path) -> list[dict]:
    """
    Parse the vehicle table from the .docx file.

    The Word document contains exactly one table with four columns:
        Sl. No | Allotted District | Vehicle No | Vehicle type

    Returns a list of dicts, one per data row (header row is skipped):
        {
            "district":  str,   # raw district name from the docx
            "reg_no":    str,   # registration number e.g. "WB-37C-7227"
            "veh_type":  str,   # raw type string e.g. "Big CDRV"
        }

    python-docx (Document) is used to open and traverse the table.
    Each row has a .cells property — a list of Cell objects — and each Cell
    has a .text property that gives all paragraph text concatenated.
    """
    # python-docx's Document class reads the .docx ZIP archive and exposes
    # paragraphs, tables, styles, etc. as Python objects.
    from docx import Document  # local import keeps the module importable even if
                                # python-docx isn't installed during unit tests

    doc = Document(str(docx_path))

    # Guard: the file must have at least one table
    if not doc.tables:
        raise ValueError(f"No tables found in {docx_path}")

    # We only care about the first (and only) table in the document
    table = doc.tables[0]

    rows_data: list[dict] = []

    # Iterate every row in the table, skipping the header (index 0)
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            # This is the header row ("Sl. No", "Allotted District", ...)
            continue

        # Extract and normalise the four cell values
        cells = [normalise(cell.text) for cell in row.cells]

        # Guard against completely empty rows (Word sometimes adds a trailing
        # blank row at the end of a table)
        if not any(cells):
            continue

        # Unpack — we don't need the serial number (cells[0]) for DB import
        _, district, reg_no, veh_type = cells

        rows_data.append({
            "district": district,
            "reg_no":   reg_no,
            "veh_type": veh_type,
        })

    return rows_data


# ─────────────────────────────────────────────────────────────────────────────
# MANAGEMENT COMMAND
# ─────────────────────────────────────────────────────────────────────────────

class Command(BaseCommand):
    """
    Django management command that seeds the Vehicle table from CDRV Details.docx.

    Inheriting from BaseCommand gives us:
      • self.stdout / self.stderr for coloured output.
      • self.style.SUCCESS / WARNING / ERROR for terminal colouring.
      • Automatic --help text generation from help = "...".
      • Standard argument parsing via add_arguments / self.options.
    """

    help = (
        "Import Civil Defence Response Vehicles from 'CDRV Details.docx' into "
        "the Vehicle table.  Uses update_or_create on registration_no so the "
        "command is safe to re-run.  Kolkata vehicles and duplicate registration "
        "numbers are skipped with a warning."
    )

    def add_arguments(self, parser):
        """
        add_arguments() is called by Django before handle() to let us declare
        custom CLI flags.  self.options[key] will hold their values in handle().
        """
        # --docx lets CI / scripts override the default file location
        parser.add_argument(
            "--docx",
            type=Path,
            default=DEFAULT_DOCX_PATH,
            help=f"Path to the CDRV Details.docx file (default: {DEFAULT_DOCX_PATH})",
        )
        # --dry-run prints what would happen without touching the database
        parser.add_argument(
            "--dry-run",
            action="store_true",
            default=False,
            help="Print actions without writing to the database.",
        )

    def handle(self, *args, **options):
        """
        handle() is the main entry point — Django calls it after parsing args.

        Flow:
          1. Parse the docx table into a list of row dicts.
          2. For each row:
             a. Resolve the district name → DB Unit (create if missing).
             b. Map the vehicle type string → VehicleType choice.
             c. Upsert the Vehicle row keyed on registration_no.
          3. Print a summary line.
        """
        docx_path: Path = options["docx"]
        dry_run: bool   = options["dry_run"]

        # ── Validate file existence ───────────────────────────────────────────
        if not docx_path.exists():
            self.stderr.write(self.style.ERROR(f"File not found: {docx_path}"))
            return

        self.stdout.write(f"Reading: {docx_path}")
        if dry_run:
            self.stdout.write(self.style.WARNING("DRY RUN — no database changes will be made."))

        # ── Parse the Word document ───────────────────────────────────────────
        rows = parse_vehicle_table(docx_path)
        self.stdout.write(f"Parsed {len(rows)} vehicle rows from the document.")

        # ── Counters for the summary ──────────────────────────────────────────
        created   = 0
        updated   = 0
        skipped   = 0

        # ── Track registration numbers already processed this run ─────────────
        # The source document has at least two duplicate reg numbers:
        #   WB-37C-6770 : Jalpaiguri (row 9) and Alipurduar (row 52)
        #   WB-41G-2767 : Darjeeling (row 4) and Kalimpong (row 50)
        # We keep the FIRST occurrence and warn on any later occurrence.
        seen_reg_nos: set[str] = set()

        # ── Process each row ──────────────────────────────────────────────────
        for row in rows:
            district = row["district"]
            reg_no   = row["reg_no"]
            veh_type = row["veh_type"]

            # ── 1. Resolve district → Unit ────────────────────────────────────
            # Check if the district name is in our mapping table
            if district not in UNIT_NAME_MAP:
                # Unknown district: print a warning and skip.
                # This should never happen once UNIT_NAME_MAP is complete,
                # but serves as a safety net for future document revisions.
                self.stdout.write(
                    self.style.WARNING(
                        f"  WARN Unknown district '{district}' — add to UNIT_NAME_MAP and re-run."
                    )
                )
                skipped += 1
                continue

            db_unit_name = UNIT_NAME_MAP[district]

            # None → intentionally skipped (e.g. Kolkata)
            if db_unit_name is None:
                self.stdout.write(
                    self.style.WARNING(
                        f"  SKIP '{district}' ({reg_no}) — district maps to None "
                        f"(needs manual unit assignment)."
                    )
                )
                skipped += 1
                continue

            # ── 2. Guard against duplicate registration numbers ───────────────
            # A registration number is unique to one vehicle; if we see the same
            # reg_no twice in the document it is a data error.  We keep the first
            # occurrence and skip all subsequent ones.
            if reg_no in seen_reg_nos:
                self.stdout.write(
                    self.style.WARNING(
                        f"  SKIP duplicate reg_no '{reg_no}' (district: '{district}') "
                        f"— already processed earlier in the document."
                    )
                )
                skipped += 1
                continue
            seen_reg_nos.add(reg_no)

            # ── 3. Resolve / create the Unit ─────────────────────────────────
            # get_or_create returns (instance, created_bool).
            # If the unit doesn't exist yet (e.g. Purba Medinipur before
            # seed_equipment was run) it is created with just a name and slug.
            if not dry_run:
                from django.utils.text import slugify
                unit, unit_created = Unit.objects.get_or_create(
                    name=db_unit_name,
                    defaults={"slug": slugify(db_unit_name)},
                )
                if unit_created:
                    self.stdout.write(
                        self.style.WARNING(f"  NEW unit created: '{db_unit_name}'")
                    )

            # ── 4. Map vehicle type string → VehicleType choice ───────────────
            # VEHICLE_TYPE_MAP covers the two known types.  Any unexpected string
            # falls back to OTHER so the import doesn't fail on document changes.
            mapped_type = VEHICLE_TYPE_MAP.get(veh_type)
            if mapped_type is None:
                self.stdout.write(
                    self.style.WARNING(
                        f"  WARN Unknown vehicle type '{veh_type}' for {reg_no} "
                        f"— defaulting to OTHER."
                    )
                )
                mapped_type = VehicleType.OTHER

            # ── 5. Dry-run path: just print what would happen ─────────────────
            if dry_run:
                self.stdout.write(
                    f"  DRY  {district!s:25s}  {reg_no:15s}  {veh_type}"
                )
                created += 1  # count as "would create" for the summary
                continue

            # ── 6. Upsert the Vehicle row ─────────────────────────────────────
            # update_or_create(lookup_kwargs, defaults=update_fields):
            #   • If a Vehicle with this registration_no exists → update its
            #     unit and vehicle_type (in case the document was corrected).
            #   • If it doesn't exist → create a new row with all fields.
            # status defaults to AVAILABLE on creation; not overwritten on update
            # so that an admin-set "MAINTENANCE" status isn't silently reset.
            vehicle, was_created = Vehicle.objects.update_or_create(
                registration_no=reg_no,
                defaults={
                    "unit":         unit,
                    "vehicle_type": mapped_type,
                    # Only set status on creation — preserve any operational
                    # status that was manually set after the initial import.
                },
            )

            # If the vehicle was just created, explicitly set AVAILABLE status
            # (update_or_create won't set it in defaults because we don't want
            # to overwrite existing status on subsequent runs).
            if was_created:
                vehicle.status = VehicleStatus.AVAILABLE
                vehicle.save(update_fields=["status"])
                created += 1
                self.stdout.write(
                    self.style.SUCCESS(
                        f"  CREATE {reg_no:15s}  {db_unit_name:25s}  {veh_type}"
                    )
                )
            else:
                updated += 1
                self.stdout.write(
                    f"  UPDATE {reg_no:15s}  {db_unit_name:25s}  {veh_type}"
                )

        # ── Summary ───────────────────────────────────────────────────────────
        if dry_run:
            self.stdout.write(
                self.style.WARNING(
                    f"\nDRY RUN complete — would process {created} vehicles "
                    f"({skipped} skipped)."
                )
            )
        else:
            self.stdout.write(
                self.style.SUCCESS(
                    f"\nDone — {created} created, {updated} updated, {skipped} skipped."
                )
            )
